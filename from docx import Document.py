from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

def create_formatted_document(filename="AutomatedDocument.docx"):
    # Create a new Document
    doc = Document()
    
    # Add a title
    title = doc.add_heading('Automated Document Creation', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add a subtitle
    subtitle = doc.add_heading('Created with Python', 1)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add some paragraphs with formatting
    p = doc.add_paragraph('This document demonstrates how to ')
    p.add_run('create and format').bold = True
    p.add_run(' Microsoft Word documents using Python code. This approach is useful for ')
    p.add_run('automated report generation').italic = True
    p.add_run(' and document processing.')
    
    # Add a section heading
    doc.add_heading('Document Sections', level=2)
    
    # Add a bullet list
    doc.add_paragraph('Here are some formatting options:', style='List Bullet')
    doc.add_paragraph('Change font styles (bold, italic, underline)', style='List Bullet')
    doc.add_paragraph('Add tables and images', style='List Bullet')
    doc.add_paragraph('Create headers and footers', style='List Bullet')
    
    # Add a section heading
    doc.add_heading('Table Example', level=2)
    
    # Add a table
    table = doc.add_table(rows=3, cols=3)
    table.style = 'Table Grid'
    
    # Add header row
    header_cells = table.rows[0].cells
    header_cells[0].text = 'Header 1'
    header_cells[1].text = 'Header 2'
    header_cells[2].text = 'Header 3'
    
    # Add data to the table
    data_cells = table.rows[1].cells
    data_cells[0].text = 'Data 1'
    data_cells[1].text = 'Data 2'
    data_cells[2].text = 'Data 3'
    
    data_cells = table.rows[2].cells
    data_cells[0].text = 'Data 4'
    data_cells[1].text = 'Data 5'
    data_cells[2].text = 'Data 6'
    
    # Add page break
    doc.add_page_break()
    
    # Add content on the new page
    doc.add_heading('Second Page', level=1)
    doc.add_paragraph('This content appears on the second page of the document.')
    
    # Add a footer
    section = doc.sections[0]
    footer = section.footer
    footer_para = footer.paragraphs[0]
    footer_para.text = "Generated using Python automation - Page "
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Save the document
    doc.save(filename)
    print(f"Document saved as {filename}")

if __name__ == "__main__":
    create_formatted_document()